# contract_templates/routes.py
from flask import render_template, send_file, current_app, abort
from flask_login import login_required, current_user
from io import BytesIO
from pathlib import Path
from urllib.parse import quote

from access_control import access_required, agent_required
from extensions import db
from models import UserProfile, utcnow

from . import contract_templates_bp

# Import python-docx la nivelul modulului
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None
    Pt = None
    Inches = None
    WD_ALIGN_PARAGRAPH = None


def _require_active_account():
    """Verifică dacă utilizatorul are abonament activ."""
    if not current_user.is_active_account():
        from flask import flash, redirect, url_for
        flash("Drepturile de contract sunt disponibile doar pentru conturi active. Activează abonamentul pentru a accesa această funcție.", "error")
        return redirect(url_for("home"))
    return None


def _get_user_profile():
    """Obține profilul utilizatorului sau datele implicite."""
    profile = UserProfile.query.filter_by(user_id=current_user.id).first()
    if profile:
        return {
            "agent_name": profile.agent_name or "",
            "agency_name": profile.agency_name or "",
            "agency_address": profile.agency_hq_address or "",
            "agency_orc": profile.agency_orc_number or "",
            "agency_cui": profile.agency_cui or "",
            "agency_iban": profile.agency_iban or "",
            "agency_bank": profile.agency_bank or "",
            "agency_administrator": profile.agency_administrator or "",
        }
    return {
        "agent_name": "",
        "agency_name": "",
        "agency_address": "",
        "agency_orc": "",
        "agency_cui": "",
        "agency_iban": "",
        "agency_bank": "",
        "agency_administrator": "",
    }


@contract_templates_bp.get("/")
@login_required
@agent_required
@access_required
def list_templates():
    """Lista tuturor draft-urilor disponibile."""
    redirect_response = _require_active_account()
    if redirect_response:
        return redirect_response
    
    templates = [
        {
            "id": "chirie",
            "title": "Fișă de vizionare - Chirie",
            "description": "Draft de fișă de vizionare pentru închiriere",
            "icon": "chirie",
        },
        {
            "id": "vanzare",
            "title": "Fișă de vizionare - Vânzare",
            "description": "Draft de fișă de vizionare pentru vânzare",
            "icon": "vanzare",
        },
        {
            "id": "prestari",
            "title": "Contract Prestări Servicii",
            "description": "Draft de contract pentru prestări servicii imobiliare",
            "icon": "prestari",
        },
        {
            "id": "inchiriere",
            "title": "Contract Închiriere",
            "description": "Draft de contract de închiriere",
            "icon": "inchiriere",
        },
        {
            "id": "exclusivitate",
            "title": "Contract Exclusivitate",
            "description": "Draft de contract de intermediere imobiliară - reprezentare exclusivă",
            "icon": "exclusivitate",
        },
    ]
    
    return render_template("contract_templates/list.html", templates=templates)


@contract_templates_bp.get("/preview/<template_id>")
@login_required
@agent_required
@access_required
def preview_template(template_id):
    """Previzualizare draft template."""
    redirect_response = _require_active_account()
    if redirect_response:
        return redirect_response
    
    profile = _get_user_profile()
    
    # Date demo pentru previzualizare
    demo_data = {
        "data_vizionarii": "15.02.2026",
        "ora_vizionarii": "14:00",
        "tip_imobil": "Apartament",
        "adresa_public": "Strada Exemplu, Nr. 1, București",
        "comision_procent": "50%",
    }
    
    if template_id == "chirie":
        return render_template(
            "contract_templates/preview_chirie.html",
            info=demo_data,
            agent_name=profile["agent_name"],
            agency_text=profile["agency_name"] or "",
        )
    elif template_id == "vanzare":
        return render_template(
            "contract_templates/preview_vanzare.html",
            info=demo_data,
            agent_name=profile["agent_name"],
            agency_text=profile["agency_name"] or "",
        )
    elif template_id == "prestari":
        return render_template(
            "contract_templates/preview_prestari.html",
            profile=profile,
        )
    elif template_id == "inchiriere":
        return render_template(
            "contract_templates/preview_inchiriere.html",
            profile=profile,
        )
    else:
        abort(404)


@contract_templates_bp.get("/download/<template_id>")
@login_required
@agent_required
@access_required
def download_template(template_id):
    """Descarcă draft template ca fișier Word (.docx)."""
    redirect_response = _require_active_account()
    if redirect_response:
        return redirect_response
    
    if not DOCX_AVAILABLE:
        from flask import flash, redirect, url_for
        flash("Biblioteca python-docx nu este instalată. Contactează suportul.", "error")
        return redirect(url_for("contract_templates.list_templates"))
    
    profile = _get_user_profile()
    
    # Creează document Word
    doc = Document()
    
    # Setări document
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4
    section.page_width = Inches(8.27)
    section.left_margin = Inches(0.98)
    section.right_margin = Inches(0.98)
    section.top_margin = Inches(0.98)
    section.bottom_margin = Inches(0.98)
    
    if template_id == "chirie":
        _build_chirie_docx(doc, profile)
        filename = "Fisa_vizionare_chirie.docx"
    elif template_id == "vanzare":
        _build_vanzare_docx(doc, profile)
        filename = "Fisa_vizionare_vanzare.docx"
    elif template_id == "prestari":
        _build_prestari_docx(doc, profile)
        filename = "Contract_prestari_servicii.docx"
    elif template_id == "inchiriere":
        _build_inchiriere_docx(doc, profile)
        filename = "Contract_inchiriere.docx"
    elif template_id == "exclusivitate":
        _build_exclusivitate_docx(doc, profile)
        filename = "Contract_exclusivitate.docx"
    else:
        abort(404)
    
    # Salvează în memorie
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    # Răspuns optimizat pentru mobile (iPhone și Android)
    response = send_file(
        file_stream,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
        conditional=False,  # Dezactivează conditional requests pentru mobile
    )
    
    # Headers pentru compatibilitate mobile (iPhone și Android)
    # Encoding corect pentru numele fișierului cu caractere speciale
    encoded_filename = quote(filename.encode('utf-8'))
    response.headers["Content-Disposition"] = f'attachment; filename="{filename}"; filename*=UTF-8\'\'{encoded_filename}'
    response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"
    response.headers["X-Content-Type-Options"] = "nosniff"
    
    return response


def _build_chirie_docx(doc, profile):
    """Construiește documentul Word pentru fișă chirie - versiune standalone, editabilă."""
    from docx.shared import Pt
    
    # Setări pentru spacing optimizat
    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)
    
    # Titlu
    title = doc.add_heading("FIȘĂ DE VIZIONARE — ÎNCHIRIERE", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph("(Draft – document pentru semnare fizică)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    
    doc.add_paragraph("Data vizionării: ___ / ___ / ______")
    doc.add_paragraph("Ora vizionării: _______")
    
    # Secțiunea 1: Părți
    doc.add_heading("1. PĂRȚI", level=2)
    doc.add_paragraph("Prestator (Agenție imobiliară):")
    doc.add_paragraph(f"Denumire: {profile['agency_name'] or '______________________________'} SRL")
    doc.add_paragraph(f"Sediu: {profile['agency_address'] or '_________________________________'}")
    doc.add_paragraph(f"CUI: {profile['agency_cui'] or '_________________________________'}")
    doc.add_paragraph("Agent imobiliar:")
    doc.add_paragraph(f"Nume și prenume: {profile['agent_name'] or '______________________________'}")
    
    # Secțiunea 2: Vizitator
    doc.add_heading("2. VIZITATOR (CLIENT)", level=2)
    doc.add_paragraph("Datele Vizitatorului se completează mai jos:")
    doc.add_paragraph("Nume și prenume: ______________________________")
    doc.add_paragraph("Telefon: ______________________________")
    doc.add_paragraph("E-mail: ______________________________")
    doc.add_paragraph("CI – Serie și număr: ______________________________")
    doc.add_paragraph("Notă: Datele din actul de identitate sunt utilizate exclusiv în scop de identificare.")
    doc.add_paragraph("Vizitatorul declară că datele furnizate sunt reale, corecte și îi aparțin. Furnizarea acestora are ca scop organizarea vizionării, desfășurarea activității de intermediere și, după caz, dovedirea intermedierii și a consimțământului exprimat prin semnare.")
    doc.add_paragraph("În situația în care anumite date sunt opționale, acestea sunt furnizate voluntar. Refuzul furnizării unor date poate limita posibilitatea de identificare în caz de contestare, fără a afecta valabilitatea prezentei fișe.")
    
    # Secțiunea 3: Imobil
    doc.add_heading("3. IMOBIL VIZIONAT", level=2)
    doc.add_paragraph("Tip imobil: ______________________________")
    doc.add_paragraph("Adresă / zonă: ______________________________")
    doc.add_paragraph("Alte detalii relevante: ______________________________")
    doc.add_paragraph("Vizitatorul confirmă că imobilul menționat mai sus, precum și informațiile relevante pentru închiriere (caracteristici, stare, condiții, disponibilitate) i-au fost prezentate prin intermediul Prestatorului, prin Agentul imobiliar.")
    doc.add_paragraph("Introducerea Vizitatorului la proprietate constituie un element esențial al activității de intermediere.")
    
    # Secțiunea 4: Obiectul fișei
    doc.add_heading("4. OBIECTUL FIȘEI", level=2)
    doc.add_paragraph("Prin prezenta, Vizitatorul confirmă efectuarea vizionării imobilului descris mai sus prin intermediul Prestatorului, în prezența Agentului imobiliar, și primirea informațiilor necesare privind condițiile de închiriere.")
    doc.add_paragraph("Prezenta fișă are rol de dovadă a intermedierii și a introducerii Vizitatorului la proprietate, fiind utilizată pentru:")
    doc.add_paragraph("• protejarea dreptului Prestatorului la comision;", style="List Bullet")
    doc.add_paragraph("• stabilirea cronologiei evenimentelor (vizionare – negociere – închiriere);", style="List Bullet")
    doc.add_paragraph("• demonstrarea legăturii dintre vizionarea realizată prin Prestator și eventuala tranzacție.", style="List Bullet")
    doc.add_paragraph("Prezenta fișă nu ține loc de contract de închiriere și nu conferă drepturi de proprietate sau folosință, însă produce efecte juridice între părți cu privire la intermediere, obligația de neeludare și plata comisionului.")
    
    # Secțiunea 5: Clauză de neeludare
    doc.add_heading("5. CLAUZĂ DE NEELUDARE A INTERMEDIERII", level=2)
    doc.add_paragraph("Vizitatorul se obligă ca, pe o perioadă de 6 (șase) luni de la data semnării prezentei fișe, să nu contacteze direct proprietarul imobilului și să nu încheie, direct sau indirect, nicio tranzacție de închiriere având ca obiect imobilul vizionat, fără participarea Prestatorului.")
    doc.add_paragraph("Obligația se aplică inclusiv în cazul:")
    doc.add_paragraph("• utilizării unor persoane interpuse (rude până la gradul IV inclusiv, prieteni, colegi, societăți controlate sau alte persoane);", style="List Bullet")
    doc.add_paragraph("• realizării tranzacției în condiții identice, similare sau modificate (preț, durată, alte clauze), dacă există legătură cu introducerea la proprietate realizată de Prestator.", style="List Bullet")
    doc.add_paragraph("Încălcarea acestei obligații atrage răspunderea contractuală a Vizitatorului. Comisionul prevăzut la secțiunea următoare reprezintă prejudiciul minim prezumat rezultat din eludarea intermedierii, fără a limita dreptul Prestatorului de a solicita repararea integrală a prejudiciului dovedit.")
    
    # Secțiunea 6: Comision
    doc.add_heading("6. COMISION", level=2)
    doc.add_paragraph("În cazul în care tranzacția de închiriere se finalizează pentru imobilul vizionat (direct sau indirect), Vizitatorul se obligă să achite Prestatorului un comision în cuantum de _______% / __________, conform înțelegerii comerciale dintre părți.")
    doc.add_paragraph("Comisionul devine exigibil la data semnării contractului de închiriere și/sau la data obținerii folosinței imobilului (de exemplu, predarea cheilor), oricare dintre acestea intervine prima.")
    doc.add_paragraph("Plata comisionului se va efectua în termen de ____ zile calendaristice de la data finalizării tranzacției. Neplata la scadență poate conduce la demersuri de recuperare și, după caz, la acțiuni legale.")
    
    # Secțiunea 7: Acord
    doc.add_heading("7. ACORD ȘI CONFIRMARE", level=2)
    doc.add_paragraph("Vizitatorul declară și confirmă că:")
    doc.add_paragraph("• a efectuat vizionarea imobilului prin intermediul Prestatorului;", style="List Bullet")
    doc.add_paragraph("• a luat cunoștință de conținutul prezentei fișe și îl acceptă integral;", style="List Bullet")
    doc.add_paragraph("• datele furnizate sunt reale și îi aparțin;", style="List Bullet")
    doc.add_paragraph("• a înțeles clauza de neeludare și obligațiile care decurg din aceasta.", style="List Bullet")
    doc.add_paragraph("Vizitatorul confirmă că a avut posibilitatea de a citi documentul înainte de semnare, că a solicitat lămuriri acolo unde a considerat necesar și că își exprimă consimțământul în mod liber.")
    
    # Secțiunea 8: GDPR
    doc.add_heading("8. PROTECȚIA DATELOR (GDPR)", level=2)
    doc.add_paragraph("Operatorul de date cu caracter personal este Prestatorul menționat la secțiunea 1. Datele sunt prelucrate în conformitate cu Regulamentul (UE) 2016/679 (GDPR) și legislația națională aplicabilă.")
    doc.add_paragraph("Scopurile prelucrării includ: organizarea vizionării, intermedierea tranzacției, comunicarea cu Vizitatorul, dovedirea intermedierii și protejarea dreptului la comision, precum și apărarea drepturilor în cazul unor litigii.")
    doc.add_paragraph("Datele nu sunt utilizate în scopuri de marketing și nu sunt vândute. Ele pot fi comunicate autorităților sau consultanților juridici doar în temei legal sau pentru apărarea unui drept.")
    doc.add_paragraph("Vizitatorul beneficiază de drepturile prevăzute de lege (acces, rectificare, opoziție, ștergere, în limitele legii).")
    
    # Secțiunea 9: Semnături
    doc.add_heading("9. SEMNĂTURI", level=2)
    doc.add_paragraph("Vizitator (Client):")
    doc.add_paragraph("Nume: ______________________________")
    doc.add_paragraph("Semnătură: __________________________")
    doc.add_paragraph("Data: ___ / ___ / ______")
    doc.add_paragraph("Agent imobiliar / Prestator:")
    doc.add_paragraph("Nume: ______________________________")
    doc.add_paragraph("Semnătură: __________________________")
    doc.add_paragraph("Data: ___ / ___ / ______")


def _build_vanzare_docx(doc, profile):
    """Construiește documentul Word pentru fișă vânzare - versiune standalone, editabilă."""
    from docx.shared import Pt
    
    # Setări pentru spacing mai generos
    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(12)
    
    # Titlu
    title = doc.add_heading("FIȘĂ DE VIZIONARE — VÂNZARE", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph("(Draft – document pentru semnare fizică)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    
    doc.add_paragraph("Data vizionării: ___ / ___ / ______")
    doc.add_paragraph("Ora vizionării: _______")
    
    # Secțiunea 1: Părți
    doc.add_heading("1. PĂRȚI", level=2)
    doc.add_paragraph("Prestator (Agenție imobiliară):")
    doc.add_paragraph(f"Denumire: {profile['agency_name'] or '______________________________'} SRL")
    doc.add_paragraph(f"Sediu: {profile['agency_address'] or '_________________________________'}")
    doc.add_paragraph(f"CUI: {profile['agency_cui'] or '_________________________________'}")
    doc.add_paragraph("Agent imobiliar:")
    doc.add_paragraph(f"Nume și prenume: {profile['agent_name'] or '______________________________'}")
    
    # Secțiunea 2: Vizitator
    doc.add_heading("2. VIZITATOR (CLIENT)", level=2)
    doc.add_paragraph("Datele Vizitatorului se completează mai jos:")
    doc.add_paragraph("Nume și prenume: ______________________________")
    doc.add_paragraph("Telefon: ______________________________")
    doc.add_paragraph("E-mail: ______________________________")
    doc.add_paragraph("CI – Serie și număr: ______________________________")
    doc.add_paragraph("Vizitatorul declară că datele furnizate sunt reale, corecte și îi aparțin. Acestea sunt utilizate exclusiv în scopul organizării vizionării, desfășurării activității de intermediere și, după caz, al dovedirii intermedierii și a consimțământului exprimat prin semnare.")
    
    # Secțiunea 3: Imobil
    doc.add_heading("3. IMOBIL VIZIONAT", level=2)
    doc.add_paragraph("Tip imobil: ______________________________")
    doc.add_paragraph("Adresă completă: ______________________________")
    doc.add_paragraph("Alte detalii relevante: ______________________________")
    doc.add_paragraph("Vizitatorul confirmă că imobilul menționat mai sus, precum și informațiile relevante pentru vânzare (caracteristici, stare, condiții, disponibilitate) i-au fost prezentate prin intermediul Prestatorului, prin Agentul imobiliar.")
    doc.add_paragraph("Introducerea Vizitatorului la proprietate constituie un element esențial al activității de intermediere.")
    
    # Secțiunea 4: Obiectul fișei
    doc.add_heading("4. OBIECTUL FIȘEI", level=2)
    doc.add_paragraph("Prin prezenta, Vizitatorul confirmă că a efectuat vizionarea imobilului descris mai sus prin intermediul Prestatorului, în prezența Agentului imobiliar, și că a primit informații privind caracteristicile, starea și condițiile de vânzare ale imobilului.")
    doc.add_paragraph("Prezenta fișă are rol de dovadă a intermedierii și a introducerii Vizitatorului la proprietate, fiind utilizată pentru protejarea dreptului Prestatorului la comision și pentru stabilirea legăturii dintre vizionare și eventuala tranzacție de vânzare-cumpărare.")
    doc.add_paragraph("Părțile înțeleg că prezenta fișă nu ține loc de antecontract sau contract de vânzare-cumpărare și nu transferă drepturi de proprietate, însă produce efecte juridice între părți cu privire la intermediere, obligația de neeludare și plata comisionului.")
    
    # Secțiunea 5: Clauză de neeludare
    doc.add_heading("5. CLAUZĂ DE NEELUDARE A INTERMEDIERII", level=2)
    doc.add_paragraph("Vizitatorul se obligă ca, pe o perioadă de 6 (șase) luni de la data semnării prezentei fișe, să nu contacteze direct proprietarul/vânzătorul imobilului și să nu încheie, direct sau indirect, nicio tranzacție de vânzare-cumpărare având ca obiect imobilul vizionat, fără participarea Prestatorului.")
    doc.add_paragraph("Această obligație se aplică inclusiv:")
    doc.add_paragraph("• prin persoane interpuse (rude până la gradul IV inclusiv, prieteni, colegi, societăți controlate sau orice alte persoane);", style="List Bullet")
    doc.add_paragraph("• în situația în care tranzacția se realizează în condiții identice, similare sau modificate (preț negociat, termene diferite, alte clauze), dacă există legătură cu introducerea la proprietate realizată de Prestator.", style="List Bullet")
    doc.add_paragraph("Încălcarea obligației de neeludare atrage răspunderea contractuală a Vizitatorului, iar comisionul prevăzut la secțiunea următoare reprezintă prejudiciul minim prezumat rezultat din eludarea intermedierii, fără a limita dreptul Prestatorului de a solicita repararea integrală a prejudiciului dovedit.")
    
    # Secțiunea 6: Comision
    doc.add_heading("6. COMISION", level=2)
    doc.add_paragraph("În cazul în care tranzacția de vânzare-cumpărare se finalizează pentru imobilul vizionat (direct sau indirect), Vizitatorul se obligă să achite Prestatorului un comision în cuantum de __________ % / __________, calculat din prețul de vânzare negociat, conform înțelegerii comerciale dintre părți.")
    doc.add_paragraph("Comisionul devine exigibil la data semnării antecontractului și/sau a contractului de vânzare-cumpărare ori la data la care tranzacția produce efecte juridice, după caz.")
    doc.add_paragraph("Plata comisionului se va efectua în termen de maximum 7 (șapte) zile calendaristice de la data finalizării tranzacției.")
    
    # Secțiunea 7: Acord
    doc.add_heading("7. ACORD ȘI CONFIRMARE", level=2)
    doc.add_paragraph("Vizitatorul declară și confirmă că:")
    doc.add_paragraph("• a efectuat vizionarea imobilului prin intermediul Prestatorului;", style="List Bullet")
    doc.add_paragraph("• a luat cunoștință de conținutul prezentei fișe și îl acceptă integral;", style="List Bullet")
    doc.add_paragraph("• datele furnizate sunt reale și îi aparțin;", style="List Bullet")
    doc.add_paragraph("• a înțeles clauza de neeludare și efectele acesteia, inclusiv obligația de plată a comisionului în caz de eludare.", style="List Bullet")
    doc.add_paragraph("Vizitatorul confirmă că a avut posibilitatea de a citi documentul înainte de semnare și că își exprimă consimțământul în mod liber.")
    
    # Secțiunea 8: Semnături
    doc.add_heading("8. SEMNĂTURI", level=2)
    doc.add_paragraph("Vizitator (Client):")
    doc.add_paragraph("Nume: ______________________________")
    doc.add_paragraph("Semnătură: __________________________")
    doc.add_paragraph("Data: ___ / ___ / ______")
    doc.add_paragraph("Agent imobiliar / Prestator:")
    doc.add_paragraph("Nume: ______________________________")
    doc.add_paragraph("Semnătură: __________________________")
    doc.add_paragraph("Data: ___ / ___ / ______")


def _build_prestari_docx(doc, profile):
    """Construiește documentul Word pentru contract prestări servicii - versiune standalone, editabilă."""
    from docx.shared import Pt
    
    # Setări pentru spacing mai generos
    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(12)
    
    title = doc.add_heading("CONTRACT DE PRESTĂRI SERVICII IMOBILIARE", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph("(Draft – document pentru completare și semnare fizică)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    
    doc.add_paragraph("Nr. ___ / Data ___ / ___ / ______")
    
    doc.add_heading("I. PĂRȚILE CONTRACTANTE", level=2)
    doc.add_heading("1.1. Prestatorul", level=3)
    doc.add_paragraph(f"Societatea {profile['agency_name'] or '______________________________'} S.R.L., persoană juridică română, cu sediul social în {profile['agency_address'] or '_________________________________'}, înregistrată la Oficiul Registrului Comerțului sub nr. {profile['agency_orc'] or '_________________________________'}, având cod unic de înregistrare {profile['agency_cui'] or '_________________________________'}, cont bancar IBAN {profile['agency_iban'] or '_________________________________'}, deschis la {profile['agency_bank'] or '_________________________________'}, reprezentată legal prin {profile['agency_administrator'] or '_________________________________'}, în calitate de ______________________________, denumită în continuare Prestatorul,")
    doc.add_paragraph("și")
    doc.add_heading("1.2. Beneficiarul", level=3)
    doc.add_paragraph("Domnul/Doamna ______________________________, identificat(ă) cu CNP ______________________________, domiciliat(ă) în ______________________________, având date de contact: telefon ______________________________ și adresă de e-mail ______________________________, denumit(ă) în continuare Beneficiarul.")
    doc.add_paragraph("Prestatorul și Beneficiarul vor fi denumiți împreună Părțile, iar individual Partea.")
    
    doc.add_heading("Art. 1. OBIECTUL CONTRACTULUI", level=2)
    doc.add_paragraph("1.1. Obiectul prezentului contract îl constituie prestarea de către Prestator a serviciilor de intermediere imobiliară, în vederea închirierii / vânzării imobilului – locuință situată la:")
    doc.add_paragraph("Adresa imobilului: ______________________________")
    doc.add_paragraph("1.2. Serviciile de intermediere includ, fără a se limita la: promovarea ofertei, organizarea și desfășurarea vizionărilor, facilitarea negocierilor dintre părți și sprijinirea Beneficiarului în vederea perfectării tranzacției.")
    
    doc.add_heading("Art. 2. DURATA CONTRACTULUI", level=2)
    doc.add_paragraph("2.1. Prezentul contract intră în vigoare la data semnării sale de către ambele Părți.")
    doc.add_paragraph("2.2. Contractul rămâne valabil până la îndeplinirea integrală și corespunzătoare a obligațiilor asumate de Părți, inclusiv plata integrală a comisionului datorat.")
    
    doc.add_heading("Art. 3. PREȚUL ȘI CONDIȚIILE DE PLATĂ", level=2)
    doc.add_paragraph("3.1. Pentru serviciile prestate în temeiul prezentului contract, Beneficiarul se obligă să achite Prestatorului un comision în cuantum de __________ lei / __________ %, stabilit de comun acord între Părți.")
    doc.add_paragraph("3.2. Comisionul devine exigibil la data finalizării tranzacției, respectiv la data semnării contractului de închiriere sau a antecontractului / contractului de vânzare-cumpărare, după caz.")
    doc.add_paragraph("3.3. Plata comisionului se va efectua în termen de ____ zile calendaristice de la data exigibilității, în baza facturii emise de Prestator.")
    
    doc.add_heading("Art. 4. OBLIGAȚIILE PĂRȚILOR", level=2)
    doc.add_heading("4.1. Obligațiile Prestatorului", level=3)
    doc.add_paragraph("a) să depună toate diligențele necesare și să acționeze cu profesionalism în vederea realizării obiectului prezentului contract;")
    doc.add_paragraph("b) să asigure desfășurarea activității de intermediere în mod loial și transparent;")
    doc.add_paragraph("c) să informeze Beneficiarul cu privire la aspectele relevante ale tranzacției.")
    doc.add_heading("4.2. Obligațiile Beneficiarului", level=3)
    doc.add_paragraph("a) să colaboreze cu Prestatorul pe durata executării contractului;")
    doc.add_paragraph("b) să furnizeze informațiile necesare realizării intermedierii;")
    doc.add_paragraph("c) să achite comisionul convenit, în condițiile prezentului contract.")
    
    doc.add_heading("Art. 5. RĂSPUNDEREA CONTRACTUALĂ ȘI FORȚA MAJORĂ", level=2)
    doc.add_paragraph("5.1. Neexecutarea sau executarea necorespunzătoare a obligațiilor asumate atrage răspunderea contractuală a Părții în culpă, în condițiile legii.")
    doc.add_paragraph("5.2. Forța majoră exonerează de răspundere Partea care o invocă, pe durata existenței acesteia, cu condiția notificării celeilalte Părți în termen de ____ zile de la apariția evenimentului și a prezentării dovezilor corespunzătoare.")
    
    doc.add_heading("Art. 6. ÎNCETAREA CONTRACTULUI", level=2)
    doc.add_paragraph("6.1. Prezentul contract încetează:")
    doc.add_paragraph("a) prin realizarea obiectului său;", style="List Bullet")
    doc.add_paragraph("b) prin acordul scris al Părților;", style="List Bullet")
    doc.add_paragraph("c) prin denunțarea unilaterală de către oricare dintre Părți, cu notificare scrisă transmisă cu ____ zile înainte.", style="List Bullet")
    doc.add_paragraph("6.2. Încetarea contractului nu afectează obligațiile născute anterior încetării, inclusiv dreptul Prestatorului la comision.")
    
    doc.add_heading("Art. 7. CONFIDENȚIALITATE ȘI PROTECȚIA DATELOR", level=2)
    doc.add_paragraph("7.1. Părțile se obligă să păstreze confidențialitatea informațiilor obținute în legătură cu prezentul contract.")
    doc.add_paragraph("7.2. Prelucrarea datelor cu caracter personal se realizează în conformitate cu legislația aplicabilă privind protecția datelor.")
    
    doc.add_heading("Art. 8. DISPOZIȚII FINALE", level=2)
    doc.add_paragraph("8.1. Prezentul contract este guvernat de legea română.")
    doc.add_paragraph("8.2. Eventualele divergențe se vor soluționa pe cale amiabilă, iar în cazul în care acest lucru nu este posibil, de către instanțele judecătorești competente.")
    
    doc.add_heading("SEMNĂTURI", level=2)
    doc.add_paragraph("BENEFICIAR")
    doc.add_paragraph("Nume: ______________________________")
    doc.add_paragraph("Semnătură: __________________________")
    doc.add_paragraph("Data: ___ / ___ / ______")
    doc.add_paragraph("PRESTATOR")
    doc.add_paragraph("Nume: ______________________________")
    doc.add_paragraph("Semnătură: __________________________")
    doc.add_paragraph("Data: ___ / ___ / ______")


def _build_inchiriere_docx(doc, profile):
    """Construiește documentul Word pentru contract închiriere - versiune standalone, editabilă."""
    from docx.shared import Pt
    
    # Setări pentru spacing mai generos
    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(12)
    
    title = doc.add_heading("CONTRACT DE ÎNCHIRIERE", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph("(Draft – pentru completare și semnare olografă)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    
    doc.add_heading("I. PĂRȚILE CONTRACTANTE", level=2)
    doc.add_heading("1.1. Proprietarul", level=3)
    doc.add_paragraph("Domnul/Doamna ______________________________, cu domiciliul în ______________________________, telefon ______________________________, adresă de e-mail ______________________________, denumit(ă) în continuare Proprietarul,")
    doc.add_paragraph("și")
    doc.add_heading("1.2. Chiriașul", level=3)
    doc.add_paragraph("Domnul/Doamna ______________________________, cu domiciliul în ______________________________, telefon ______________________________, adresă de e-mail ______________________________, denumit(ă) în continuare Chiriașul.")
    doc.add_paragraph("Proprietarul și Chiriașul vor fi denumiți împreună Părțile.")
    
    doc.add_heading("Art. 2. OBIECTUL CONTRACTULUI. DESCRIEREA IMOBILULUI", level=2)
    doc.add_paragraph("2.1. Proprietarul dă în folosință Chiriașului imobilul ce face obiectul prezentului contract, în scop exclusiv locativ.")
    doc.add_paragraph("2.2. Imobilul este situat la adresa: ______________________________")
    doc.add_paragraph("2.3. Proprietarul declară pe propria răspundere că este titularul dreptului de proprietate asupra imobilului, că acesta nu este scos din circuitul civil și nu este grevat de sarcini sau drepturi reale care ar putea afecta ori împiedica închirierea sa.")
    
    doc.add_heading("Art. 3. DOTĂRI. PREDAREA-PRIMIREA IMOBILULUI", level=2)
    doc.add_paragraph("3.1. Predarea imobilului se realizează în stare corespunzătoare folosinței, cu dotările și bunurile existente la data predării.")
    doc.add_paragraph("3.2. Părțile pot conveni întocmarea unui proces-verbal de predare-primire, care va reflecta starea imobilului și dotările existente la momentul predării.")
    
    doc.add_heading("Art. 4. SUBÎNCHIRIEREA ȘI CEDAREA FOLOSINȚEI", level=2)
    doc.add_paragraph("4.1. Chiriașului îi este interzis să subînchirieze, să cedeze sau să transmită, total ori parțial, folosința imobilului, sub orice formă, fără acordul scris, prealabil, al Proprietarului.")
    
    doc.add_heading("Art. 5. DURATA CONTRACTULUI", level=2)
    doc.add_paragraph("5.1. Prezentul contract se încheie pe durata convenită de Părți.")
    doc.add_paragraph("5.2. Contractul poate fi prelungit prin acordul scris al Părților, consemnat într-un act adițional.")
    
    doc.add_heading("Art. 6. CHIRIA, GARANȚIA ȘI MODALITATEA DE PLATĂ", level=2)
    doc.add_paragraph("6.1. Chiria lunară este stabilită de Părți la suma de ______________________________.")
    doc.add_paragraph("6.2. Părțile pot conveni constituirea unei garanții (depozit) în vederea acoperirii eventualelor prejudicii și/sau obligații restante ale Chiriașului.")
    doc.add_paragraph("6.3. La semnarea prezentului contract, Părțile pot consemna sumele achitate cu titlu de chirie și/sau garanție, după caz.")
    doc.add_paragraph("6.4. Garanția se restituie Chiriașului la încetarea contractului, după predarea imobilului și stingerea tuturor obligațiilor contractuale (chirie, utilități, eventuale daune).")
    doc.add_paragraph("6.5. În situația constatării unor prejudicii, Proprietarul are dreptul de a reține din garanție contravaloarea acestora, justificată prin documente și/sau constatări efectuate de comun acord.")
    doc.add_paragraph("6.6. Chiria se achită lunar, la data convenită de Părți.")
    
    doc.add_heading("Art. 7. UTILITĂȚI ȘI ALTE CHELTUIELI", level=2)
    doc.add_paragraph("7.1. Chiriașul suportă costurile utilităților și ale serviciilor aferente folosinței imobilului, inclusiv, dar fără a se limita la: apă, energie electrică, gaze naturale, internet/cablu, întreținere, salubritate, în baza facturilor sau documentelor justificative.")
    
    doc.add_heading("Art. 8. DREPTURI ȘI OBLIGAȚII PRIVIND FOLOSINȚA IMOBILULUI", level=2)
    doc.add_paragraph("8.1. Chiriașul se obligă să folosească imobilul cu prudență și diligență, respectând destinația locativă, normele de conviețuire socială și legislația aplicabilă.")
    doc.add_paragraph("8.2. Chiriașul se obligă să mențină curățenia în imobil, să respecte orele legale de liniște și să nu efectueze modificări sau amenajări fără acordul scris al Proprietarului.")
    doc.add_paragraph("8.3. Este interzisă desfășurarea de activități comerciale, stabilirea sediului social sau a unui punct de lucru în imobil, fără acordul scris al Proprietarului.")
    
    doc.add_heading("Art. 9. ACCESUL PROPRIETARULUI ÎN IMOBIL", level=2)
    doc.add_paragraph("9.1. Proprietarul are dreptul de a avea acces în imobil, cu notificarea prealabilă a Chiriașului, pentru verificări sau intervenții necesare.")
    doc.add_paragraph("9.2. În caz de urgență (avarii, incendii, inundații sau alte situații ce pot produce prejudicii), Proprietarul poate avea acces în imobil fără notificare prealabilă.")
    
    doc.add_heading("Art. 10. NEPLATA CHIRIEI", level=2)
    doc.add_paragraph("10.1. În cazul întârzierii la plata chiriei, Proprietarul are dreptul de a solicita penalități de întârziere, conform înțelegerii Părților sau legislației aplicabile.")
    doc.add_paragraph("10.2. Proprietarul are dreptul de a reține din garanție sumele datorate cu titlu de chirie restantă, utilități neachitate sau alte prejudicii cauzate de Chiriaș.")
    
    doc.add_heading("Art. 11. ÎNCETAREA CONTRACTULUI. PREAVIZ", level=2)
    doc.add_paragraph("11.1. Contractul poate fi denunțat unilateral de oricare dintre Părți, cu respectarea unui termen de preaviz stabilit de comun acord, prin notificare scrisă.")
    doc.add_paragraph("11.2. La încetarea contractului, Chiriașul va preda imobilul în starea în care l-a primit, cu uzura normală aferentă folosinței.")
    
    doc.add_heading("Art. 12. REZILIEREA CONTRACTULUI PENTRU CULPĂ", level=2)
    doc.add_paragraph("12.1. Proprietarul poate rezilia prezentul contract, fără acordarea unui termen de preaviz, în cazul încălcării grave a obligațiilor contractuale de către Chiriaș, inclusiv, dar fără a se limita la:")
    doc.add_paragraph("a) neplata chiriei și/sau a utilităților;", style="List Bullet")
    doc.add_paragraph("b) producerea de distrugeri sau deteriorări semnificative ale imobilului;", style="List Bullet")
    doc.add_paragraph("c) subînchirierea sau cedarea folosinței fără acordul scris al Proprietarului;", style="List Bullet")
    doc.add_paragraph("d) desfășurarea de activități ilegale sau contrare destinației imobilului.", style="List Bullet")
    doc.add_paragraph("12.2. Rezilierea produce efecte de la data comunicării notificării scrise către Chiriaș.")
    
    doc.add_heading("Art. 13. FORȚA MAJORĂ", level=2)
    doc.add_paragraph("13.1. Niciuna dintre Părți nu răspunde pentru neexecutarea sau executarea cu întârziere a obligațiilor contractuale, dacă aceasta este cauzată de un eveniment de forță majoră, astfel cum este definit de lege.")
    doc.add_paragraph("13.2. Partea care invocă forța majoră are obligația de a notifica cealaltă Parte în cel mai scurt timp și de a lua toate măsurile rezonabile pentru limitarea efectelor evenimentului.")
    
    doc.add_heading("Art. 14. PRELUCRAREA DATELOR CU CARACTER PERSONAL", level=2)
    doc.add_paragraph("14.1. Părțile declară că au luat la cunoștință faptul că datele cu caracter personal furnizate în legătură cu prezentul contract vor fi prelucrate exclusiv în scopul executării obligațiilor contractuale și al îndeplinirii obligațiilor legale, în conformitate cu Regulamentul (UE) 2016/679 (GDPR) și legislația națională aplicabilă.")
    
    doc.add_heading("SEMNĂTURI", level=2)
    doc.add_paragraph("PROPRIETAR")
    doc.add_paragraph("Nume: ______________________________")
    doc.add_paragraph("Semnătură: __________________________")
    doc.add_paragraph("CHIRIAȘ")
    doc.add_paragraph("Nume: ______________________________")
    doc.add_paragraph("Semnătură: __________________________")


def _build_exclusivitate_docx(doc, profile):
    """Construiește documentul Word pentru contract exclusivitate - versiune standalone, editabilă."""
    from docx.shared import Pt
    
    # Setări pentru spacing mai generos
    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(12)
    
    title = doc.add_heading("CONTRACT DE INTERMEDIERE IMOBILIARĂ", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph("(Reprezentare exclusivă a Clientului Vânzător)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    
    p2 = doc.add_paragraph("(Draft – pentru completare și semnare olografă)")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(12)
    
    doc.add_paragraph("Contract nr. __________")
    doc.add_paragraph("Data semnării: ___ / ___ / ______ (ziua și ora)")
    
    doc.add_heading("I. PĂRȚILE CONTRACTANTE", level=2)
    doc.add_heading("1.1. Agenția", level=3)
    doc.add_paragraph(f"Societatea {profile['agency_name'] or '______________________________'} S.R.L., cu sediul în {profile['agency_address'] or '_________________________________'}, înregistrată la Oficiul Registrului Comerțului sub nr. {profile['agency_orc'] or '_________________________________'}, având C.U.I. {profile['agency_cui'] or '_________________________________'}, reprezentată legal prin {profile['agency_administrator'] or '_________________________________'}, în calitate de ______________________________, denumită în continuare Agenția, în calitate de intermediar exclusiv,")
    doc.add_paragraph("și")
    doc.add_heading("1.2. Clientul", level=3)
    doc.add_paragraph("Domnul/Doamna / Societatea ______________________________, cu domiciliul / sediul în ______________________________, identificat(ă) prin ______________________________, denumit(ă) în continuare Clientul.")
    doc.add_paragraph("Agenția și Clientul vor fi denumiți împreună Părțile.")
    
    doc.add_heading("Art. 2. OBIECTUL CONTRACTULUI. OFERTA IMOBILIARĂ. REPREZENTARE EXCLUSIVĂ", level=2)
    doc.add_paragraph("2.1. Clientul intenționează să încheie o tranzacție imobiliară având ca obiect înstrăinarea dreptului de proprietate asupra imobilului descris în prezentul contract și în Anexa 1 – Oferta Imobiliară, denumită în continuare Oferta.")
    doc.add_paragraph("2.2. În acest scop, Clientul mandatează Agenția să efectueze, în numele și în interesul său, servicii de promovare, consultanță și intermediere imobiliară, în schimbul unui comision, în condițiile prezentului contract.")
    doc.add_paragraph("2.3. Prin tranzacție imobiliară, Părțile înțeleg orice act sau operațiune juridică ce are ca efect transferul dreptului de proprietate asupra imobilului (inclusiv, dar fără a se limita la: contract de vânzare-cumpărare, schimb, hotărâre judecătorească ce ține loc de contract, adjudecare etc.).")
    doc.add_paragraph("2.4. Clientul desemnează Agenția unic reprezentant pentru promovarea și intermedierea Ofertei pe durata prezentului contract.")
    doc.add_paragraph("2.5. Prețul de listare acceptat de Client la data semnării prezentului contract este de ______________________________ euro, acesta putând fi modificat de Client cu informarea prealabilă a Agenției.")
    
    doc.add_heading("Art. 3. DURATA CONTRACTULUI", level=2)
    doc.add_paragraph("3.1. Prezentul contract se încheie pe o durată de ______________________________ luni, începând cu data semnării.")
    
    doc.add_heading("Art. 4. COMISIONUL ȘI CONDIȚIILE DE PLATĂ", level=2)
    doc.add_paragraph("4.1. Comisionul datorat de Client Agenției este de ______________________________ % din valoarea totală a tranzacției imobiliare.")
    doc.add_paragraph("4.2. Comisionul este scadent și datorat în următoarele situații:")
    doc.add_paragraph("a) dacă pe durata contractului Imobilul face obiectul unei tranzacții imobiliare, indiferent dacă terțul a fost identificat de Agenție sau nu;", style="List Bullet")
    doc.add_paragraph("b) dacă tranzacția se încheie în termen de 12 luni de la încetarea contractului cu un terț introdus de Agenție pe durata contractului.", style="List Bullet")
    doc.add_paragraph("4.3. Plata comisionului se efectuează:")
    doc.add_paragraph("a) în termen de ___ zile de la semnarea actului de înstrăinare;", style="List Bullet")
    doc.add_paragraph("b) prin virament bancar;", style="List Bullet")
    doc.add_paragraph("c) în echivalent lei, la cursul BNR din data semnării actului de înstrăinare.", style="List Bullet")
    doc.add_paragraph("4.4. În cazul încasării unui avans din prețul tranzacției, Clientul va achita Agenției un avans de ___% din comision, în termen de ___ zile lucrătoare.")
    
    doc.add_heading("Art. 5. DREPTURILE ȘI OBLIGAȚIILE PĂRȚILOR", level=2)
    doc.add_heading("5.1. Drepturile și obligațiile Agenției", level=3)
    doc.add_paragraph("Agenția se obligă, fără a se limita la:")
    doc.add_paragraph("a) elaborarea și implementarea unui plan de promovare;", style="List Bullet")
    doc.add_paragraph("b) promovarea Ofertei prin canale specifice;", style="List Bullet")
    doc.add_paragraph("c) organizarea vizionărilor;", style="List Bullet")
    doc.add_paragraph("d) facilitarea negocierilor;", style="List Bullet")
    doc.add_paragraph("e) informarea Clientului cu privire la terți interesați;", style="List Bullet")
    doc.add_paragraph("f) participarea la întâlniri și negocieri, la solicitarea Clientului.", style="List Bullet")
    doc.add_heading("5.2. Drepturile și obligațiile Clientului", level=3)
    doc.add_paragraph("Clientul se obligă, fără a se limita la:")
    doc.add_paragraph("a) furnizarea de informații reale și complete privind imobilul;", style="List Bullet")
    doc.add_paragraph("b) colaborarea exclusivă cu Agenția;", style="List Bullet")
    doc.add_paragraph("c) neacordarea mandatului altor intermediari;", style="List Bullet")
    doc.add_paragraph("d) achitarea comisionului conform contractului;", style="List Bullet")
    doc.add_paragraph("e) comunicarea oricărei negocieri sau tranzacții.", style="List Bullet")
    
    doc.add_heading("Art. 6. RĂSPUNDEREA CONTRACTUALĂ", level=2)
    doc.add_paragraph("6.1. Neexecutarea sau executarea necorespunzătoare a obligațiilor asumate atrage răspunderea contractuală a părții în culpă.")
    doc.add_paragraph("6.2. Întârzierea la plată atrage penalități de ___% pe zi de întârziere, fără a fi necesară punerea în întârziere.")
    doc.add_paragraph("6.3. Agenția este îndreptățită la despăgubiri egale cu dublul comisionului în cazurile de fraudare a intereselor sale.")
    
    doc.add_heading("Art. 7. ÎNCETAREA CONTRACTULUI", level=2)
    doc.add_paragraph("7.1. Contractul încetează:")
    doc.add_paragraph("a) prin ajungerea la termen;", style="List Bullet")
    doc.add_paragraph("b) prin acordul scris al Părților;", style="List Bullet")
    doc.add_paragraph("c) prin denunțare unilaterală, în condițiile prezentului contract;", style="List Bullet")
    doc.add_paragraph("d) prin reziliere pentru culpă.", style="List Bullet")
    
    doc.add_heading("Art. 8. FORȚA MAJORĂ", level=2)
    doc.add_paragraph("8.1. Forța majoră exonerează de răspundere partea care o invocă, în condițiile legii.")
    
    doc.add_heading("Art. 9. DENUNȚAREA UNILATERALĂ", level=2)
    doc.add_paragraph("9.1. Clientul poate denunța unilateral contractul cu notificare scrisă transmisă Agenției cu ___ zile înainte, cu respectarea condițiilor stabilite în prezentul contract.")
    doc.add_paragraph("9.2. Agenția poate denunța unilateral contractul cu un preaviz de ___ zile.")
    
    doc.add_heading("Art. 10. DISPOZIȚII FINALE", level=2)
    doc.add_paragraph("10.1. Prezentul contract este guvernat de legea română.")
    doc.add_paragraph("10.2. Litigiile se soluționează de instanțele competente de la locul situării imobilului.")
    doc.add_paragraph("10.3. Contractul reprezintă voința Părților și înlătură orice altă înțelegere anterioară.")
    
    doc.add_heading("SEMNĂTURI", level=2)
    doc.add_paragraph("CLIENT")
    doc.add_paragraph("Nume: ______________________________")
    doc.add_paragraph("Semnătură: __________________________")
    doc.add_paragraph("AGENȚIE")
    doc.add_paragraph("Nume: ______________________________")
    doc.add_paragraph("Semnătură: __________________________")
    
    doc.add_heading("ANEXA 1 – OFERTA IMOBILIARĂ", level=2)
    doc.add_paragraph("(Draft – de completat)")
    doc.add_paragraph("Adresă imobil: ___________________________________________")
    doc.add_paragraph("[Structura chestionarului rămâne identică, dar fără completări]")
    
    doc.add_heading("ANEXA 2 – TERȚI INTRODUȘI", level=2)
    doc.add_paragraph("(Draft tabel gol)")
    doc.add_paragraph("Nr. crt.\tNume terț\tCNP / CUI\tData introducerii\tSemnătură Client")
